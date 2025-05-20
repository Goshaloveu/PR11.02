from PyQt6.QtCore import Qt, pyqtSignal, pyqtSlot, QSize
from PyQt6.QtGui import QIcon
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QScrollArea,
    QFormLayout, QDialog, QMessageBox, QFileDialog, QCheckBox
)

from qfluentwidgets import (
    FluentIcon, InfoBar, Dialog, ComboBox, ScrollArea,
    PushButton, LineEdit, SpinBox, MessageBox, IndeterminateProgressBar,
    StrongBodyLabel, BodyLabel, CaptionLabel, SubtitleLabel, CardWidget,
    ExpandLayout, SimpleCardWidget, ToggleButton, InfoBarPosition
)

from ...common.db.controller import ProviderController, MaterialController
from ...common.db.models_pydantic import ProviderCreate, ProviderUpdate, Material
from ...common.db.database import SessionLocal
from ...common.signal_bus import signalBus
import uuid
from datetime import datetime
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import re

class SuppliersInterface(QWidget):
    def __init__(self, user_data, parent=None):
        super().__init__(parent=parent)
        self.user_data = user_data
        
        # Initialize controllers
        self.provider_controller = ProviderController()
        self.material_controller = MaterialController()
        
        # Setup UI
        self._setup_ui()
        
        # Load suppliers
        self.load_suppliers()
        
    def _setup_ui(self):
        # Main layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        
        # Header layout
        header_layout = QHBoxLayout()
        
        # Title
        title_label = SubtitleLabel("Список поставщиков")
        header_layout.addWidget(title_label)
        
        header_layout.addStretch()
        
        # Add supplier button
        self.add_button = PushButton("Добавить поставщика")
        self.add_button.setIcon(FluentIcon.ADD)
        self.add_button.clicked.connect(self.add_supplier)
        header_layout.addWidget(self.add_button)
        
        layout.addLayout(header_layout)
        
        # Scroll area for supplier cards
        self.scroll_widget = QWidget()
        # Remove fixed width constraint that's causing scrolling issues
        # self.scroll_widget.setMaximumWidth(1200)
        
        # Use QVBoxLayout for better vertical scrolling instead of ExpandLayout
        self.scroll_layout = QVBoxLayout(self.scroll_widget)
        self.scroll_layout.setSpacing(10)
        self.scroll_layout.setContentsMargins(5, 5, 5, 5)
        
        self.scroll_area = ScrollArea(self)
        self.scroll_area.setWidget(self.scroll_widget)
        self.scroll_area.setWidgetResizable(True)
        # Ensure horizontal scrollbar never appears
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        # Ensure vertical scrollbar appears when needed
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        layout.addWidget(self.scroll_area, 1)  # Add stretch factor to take available space
        
        # Progress bar for loading
        self.progress_bar = IndeterminateProgressBar(self)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
    def load_suppliers(self):
        """Load suppliers from database"""
        self.progress_bar.setVisible(True)
        
        # Clear existing cards
        while self.scroll_layout.count():
            item = self.scroll_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        try:
            db = SessionLocal()
            
            # Get all suppliers
            suppliers = self.provider_controller.get_all(db)
            
            if not suppliers:
                # No suppliers message
                no_suppliers_label = BodyLabel("Нет поставщиков. Добавьте первого поставщика с помощью кнопки выше.")
                self.scroll_layout.addWidget(no_suppliers_label)
                return
                
            # Create cards for each supplier
            for supplier in suppliers:
                # Create card
                card = SupplierCard(supplier, self)
                card.edit_clicked.connect(self.edit_supplier)
                card.delete_clicked.connect(self.delete_supplier)
                card.request_clicked.connect(self.generate_request)
                
                self.scroll_layout.addWidget(card)
            
        except Exception as e:
            InfoBar.error(
                title="Ошибка",
                content=f"Не удалось загрузить поставщиков: {str(e)}",
                parent=self
            )
        finally:
            self.progress_bar.setVisible(False)
            
    def add_supplier(self):
        """Add new supplier"""
        dialog = SupplierDialog(parent=self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_suppliers()
            # Добавляем всплывающее сообщение об успешном создании
            InfoBar.success(
                title="Успех",
                content="Новый поставщик успешно добавлен",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
        else:
            # Если пользователь отменил операцию
            InfoBar.info(
                title="Информация",
                content="Добавление поставщика отменено",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            
    def edit_supplier(self, supplier_id):
        """Edit supplier"""
        dialog = SupplierDialog(supplier_id=supplier_id, parent=self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_suppliers()
            # Добавляем всплывающее сообщение об успешном редактировании
            InfoBar.success(
                title="Успех",
                content=f"Поставщик успешно отредактирован",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
        else:
            # Если пользователь отменил операцию или произошла ошибка
            InfoBar.info(
                title="Информация",
                content="Редактирование поставщика отменено",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            
    def delete_supplier(self, supplier_id):
        """Delete supplier after confirmation"""
        
        try:
            db = SessionLocal()
            
            # Получаем имя поставщика для более информативного сообщения
            supplier = self.provider_controller.get_one(db, supplier_id)
            if not supplier:
                InfoBar.error(
                    title="Ошибка",
                    content="Поставщик не найден",
                    parent=self,
                    position=InfoBarPosition.TOP,
                    duration=3000
                )
                return
                
            supplier_name = supplier.name
            
            confirm = MessageBox(
                "Подтверждение удаления",
                f"Вы уверены, что хотите удалить поставщика '{supplier_name}'? Это действие нельзя отменить.",
                self
            )
            
            if confirm.exec():
                result = self.provider_controller.delete(db, supplier_id)
                if result:
                    self.load_suppliers()
                    InfoBar.success(
                        title="Успех",
                        content=f"Поставщик '{supplier_name}' успешно удален",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
                else:
                    InfoBar.error(
                        title="Ошибка",
                        content=f"Не удалось удалить поставщика '{supplier_name}'",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
            else:
                # Если пользователь отменил удаление
                InfoBar.info(
                    title="Информация",
                    content=f"Удаление поставщика '{supplier_name}' отменено",
                    parent=self,
                    position=InfoBarPosition.TOP,
                    duration=3000
                )
        except Exception as e:
            InfoBar.error(
                title="Ошибка",
                content=f"Не удалось удалить поставщика: {str(e)}",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
        finally:
            db.close()
                
    def generate_request(self, supplier_id):
        """Generate material request document"""
        dialog = MaterialRequestDialog(supplier_id, self)
        result = dialog.exec()
        # Добавляем всплывающее сообщение о результате операции
        if result == QDialog.DialogCode.Accepted:
            InfoBar.success(
                title="Запрос материалов",
                content="Запрос материалов успешно сформирован и открыт для просмотра",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
        else:
            InfoBar.info(
                title="Информация",
                content="Формирование запроса материалов отменено",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )


class SupplierCard(CardWidget):
    edit_clicked = pyqtSignal(str)
    delete_clicked = pyqtSignal(str)
    request_clicked = pyqtSignal(str)
    
    def __init__(self, supplier, parent=None):
        super().__init__(parent)
        self.supplier = supplier
        self.setFixedHeight(180)
        # Don't set fixed width to allow responsive layout
        self.setMinimumWidth(500)  # Minimum width for readability
        
        # Setup layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 15, 20, 15)
        layout.setSpacing(10)
        
        # Supplier name
        name_label = StrongBodyLabel(supplier.name)
        layout.addWidget(name_label)
        
        # Supplier details
        details_layout = QFormLayout()
        details_layout.setHorizontalSpacing(10)
        details_layout.setVerticalSpacing(8)
        details_layout.setLabelAlignment(Qt.AlignmentFlag.AlignRight)
        
        # INN
        inn_label = BodyLabel(f"{supplier.inn}")
        details_layout.addRow(BodyLabel("ИНН:"), inn_label)
        
        # Phone
        if supplier.phone:
            # Отображаем телефон с префиксом +7
            phone_label = BodyLabel(f"+7{supplier.phone}")
            details_layout.addRow(BodyLabel("Телефон:"), phone_label)
        
        # Email
        if supplier.mail:
            email_label = BodyLabel(supplier.mail)
            details_layout.addRow(BodyLabel("Email:"), email_label)
        
        # Address
        if supplier.address:
            address_label = BodyLabel(supplier.address)
            details_layout.addRow(BodyLabel("Адрес:"), address_label)
            
        layout.addLayout(details_layout)
        
        # Actions layout
        actions_layout = QHBoxLayout()
        actions_layout.setContentsMargins(0, 10, 0, 0)
        actions_layout.setSpacing(10)
        
        # Edit button - Fix button connections by using lambdas with explicit capture
        edit_button = PushButton("Редактировать")
        edit_button.setIcon(FluentIcon.EDIT)
        # Use a different approach to connect signals for stability
        edit_button.clicked.connect(self._on_edit_clicked)
        actions_layout.addWidget(edit_button)
        
        # Request materials button
        request_button = PushButton("Запрос материалов")
        request_button.setIcon(FluentIcon.DOCUMENT)
        # Use a different approach to connect signals for stability
        request_button.clicked.connect(self._on_request_clicked)
        actions_layout.addWidget(request_button)
        
        # Delete button
        delete_button = PushButton("Удалить")
        delete_button.setIcon(FluentIcon.DELETE)
        # Use a different approach to connect signals for stability
        delete_button.clicked.connect(self._on_delete_clicked)
        actions_layout.addWidget(delete_button)
        
        layout.addLayout(actions_layout)

    # Add helper methods to emit signals with correct supplier ID
    def _on_edit_clicked(self):
        self.edit_clicked.emit(self.supplier.id)
        
    def _on_request_clicked(self):
        self.request_clicked.emit(self.supplier.id)
        
    def _on_delete_clicked(self):
        self.delete_clicked.emit(self.supplier.id)


class SupplierDialog(MessageBox):
    def __init__(self, supplier_id=None, parent=None):
        # Set dialog title
        self.supplier_id = supplier_id
        title = "Редактирование поставщика" if supplier_id else "Добавление поставщика"
        
        # Initialize with title and placeholder content
        super().__init__(
            title=title,
            content="",
            parent=parent
        )
        
        # Replace content with our widget
        self.contentLabel.hide()
        
        # Create custom content
        self.form_widget = QWidget()
        form_layout = QVBoxLayout(self.form_widget)
        form_layout.setSpacing(15)
        
        # Form layout
        fields_layout = QFormLayout()
        fields_layout.setVerticalSpacing(10)
        fields_layout.setHorizontalSpacing(15)
        
        # Name
        self.name_edit = LineEdit()
        self.name_edit.setPlaceholderText("Введите название поставщика")
        fields_layout.addRow("Название*:", self.name_edit)
        
        # INN
        self.inn_edit = LineEdit()
        self.inn_edit.setPlaceholderText("Введите ИНН (10-12 цифр)")
        fields_layout.addRow("ИНН*:", self.inn_edit)
        
        # Phone
        self.phone_edit = LineEdit()
        self.phone_edit.setPlaceholderText("Введите телефон (формат: +7XXXXXXXXXX)")
        fields_layout.addRow("Телефон:", self.phone_edit)
        # При создании нового поставщика сразу устанавливаем префикс +7
        if not supplier_id:
            self.phone_edit.setText("+7")
        
        # Email
        self.email_edit = LineEdit()
        self.email_edit.setPlaceholderText("Введите email")
        fields_layout.addRow("Email:", self.email_edit)
        
        # Address
        self.address_edit = LineEdit()
        self.address_edit.setPlaceholderText("Введите адрес")
        fields_layout.addRow("Адрес:", self.address_edit)
        
        form_layout.addLayout(fields_layout)
        
        # Required fields note
        note_label = CaptionLabel("* - обязательные поля")
        form_layout.addWidget(note_label)
        
        # Insert our form in place of content label
        self.textLayout.insertWidget(2, self.form_widget)
        
        # Update dialog size
        self.widget.setMinimumWidth(500)
        self.widget.setMinimumHeight(400)
        
        # Update button text
        self.yesButton.setText("Сохранить")
        self.cancelButton.setText("Отмена")
        
        # Change default behavior 
        self.yesButton.clicked.disconnect()
        self.yesButton.clicked.connect(self.save_supplier)
        
        # Initialize controller
        self.provider_controller = ProviderController()
        
        # Load supplier data if editing
        if supplier_id:
            self.load_supplier_data()
        
    def load_supplier_data(self):
        """Load supplier data for editing"""
        try:
            db = SessionLocal()
            
            # Get supplier
            supplier = self.provider_controller.get_one(db, self.supplier_id)
            if not supplier:
                InfoBar.error(
                    title="Ошибка",
                    content="Поставщик не найден",
                    parent=self
                )
                self.reject()
                return
                
            # Set form values
            self.name_edit.setText(supplier.name)
            self.inn_edit.setText(supplier.inn)
            
            # Обработка телефона - всегда добавляем +7 к номеру для отображения
            current_phone = supplier.phone if supplier.phone else ""
            if current_phone:
                # В БД хранятся только цифры без префикса, для отображения добавляем +7
                self.phone_edit.setText(f"+7{current_phone}")
            else:
                # По умолчанию устанавливаем +7
                self.phone_edit.setText("+7")
                
            if supplier.mail:
                self.email_edit.setText(supplier.mail)
                
            if supplier.address:
                self.address_edit.setText(supplier.address)
                
        except Exception as e:
            InfoBar.error(
                title="Ошибка",
                content=f"Не удалось загрузить данные поставщика: {str(e)}",
                parent=self
            )
            self.reject()
        finally:
            db.close()
            
    def validate_form(self):
        """Validate form data"""
        # Required fields
        if not self.name_edit.text().strip():
            InfoBar.error(
                title="Ошибка",
                content="Название поставщика обязательно для заполнения",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            self.name_edit.setFocus()
            return False
            
        if not self.inn_edit.text().strip():
            InfoBar.error(
                title="Ошибка",
                content="ИНН поставщика обязателен для заполнения",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            self.inn_edit.setFocus()
            return False
            
        # INN format (10-12 digits)
        inn = self.inn_edit.text().strip()
        if not inn.isdigit() or len(inn) < 10 or len(inn) > 12:
            InfoBar.error(
                title="Ошибка",
                content="ИНН должен содержать 10-12 цифр",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            self.inn_edit.setFocus()
            return False
            
        # Phone format (if provided)
        phone = self.phone_edit.text().strip()
        # if phone and not (phone.startswith("+7") and len(phone) == 12) and not (phone.startswith("8") and len(phone) == 11):
        #     InfoBar.error(
        #         title="Ошибка",
        #         content="Телефон должен быть в формате +7XXXXXXXXXX или 8XXXXXXXXXX",
        #         parent=self,
        #         position=InfoBarPosition.TOP,
        #         duration=3000
        #     )
        #     self.phone_edit.setFocus()
        #     return False
            
        # Email format (if provided)
        email = self.email_edit.text().strip()
        if email and "@" not in email:
            InfoBar.error(
                title="Ошибка",
                content="Некорректный формат email",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            self.email_edit.setFocus()
            return False
            
        return True
        
    def save_supplier(self):
        """Save supplier data"""
        if not self.validate_form():
            return
            
        try:
            db = SessionLocal()
            
            # Prepare data
            name = self.name_edit.text().strip()
            inn = self.inn_edit.text().strip()
            
            # Минимальная обработка телефона для соответствия регулярному выражению в базе данных
            phone = self.phone_edit.text().strip()
            
            # Если телефон пустой, передаем None
            if not phone or phone == '+':
                phone = None
                print("DEBUG: Phone is empty or just '+', setting to None")
            else:
                # Оставляем только цифры, удаляем все остальные символы, включая '+'
                phone_digits = re.sub(r'[^0-9]', '', phone)
                
                # Если после очистки телефон пуст, передаем None
                if not phone_digits:
                    phone = None
                    print("DEBUG: Phone after cleaning is empty, setting to None")
                else:
                    # Если номер начинается с 7 или 8, берем последние 10 цифр
                    if phone_digits.startswith('7') or phone_digits.startswith('8'):
                        phone = phone_digits[-10:]
                    # Если длина 10, просто используем эти 10 цифр
                    elif len(phone_digits) == 10:
                        phone = phone_digits
                    # В других случаях берем последние 10 цифр, если их достаточно
                    elif len(phone_digits) > 10:
                        phone = phone_digits[-10:]
                    else:
                        phone = phone_digits
                    
                    print(f"DEBUG: Extracted phone digits: '{phone}'")
            
            email = self.email_edit.text().strip() or None
            address = self.address_edit.text().strip() or None
            
            print(f"DEBUG: Final values - name: '{name}', inn: '{inn}', phone: '{phone}', email: '{email}', address: '{address}'")
            
            if self.supplier_id:
                # Update existing supplier
                print(f"DEBUG: Updating supplier with ID: {self.supplier_id}")
                update_data = ProviderUpdate(
                    name=name,
                    inn=inn,
                    phone=phone,
                    mail=email,
                    address=address
                )
                
                result = self.provider_controller.update(db, self.supplier_id, update_data)
                
                if result:
                    InfoBar.success(
                        title="Успех",
                        content=f"Поставщик '{name}' успешно обновлен",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
                    self.accept()
                else:
                    InfoBar.error(
                        title="Ошибка",
                        content=f"Не удалось обновить поставщика '{name}'",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
            else:
                # Create new supplier
                create_data = ProviderCreate(
                    id=str(uuid.uuid4()),
                    name=name,
                    inn=inn,
                    phone=phone,
                    mail=email,
                    address=address
                )
                
                result = self.provider_controller.create(db, create_data)
                
                if result:
                    InfoBar.success(
                        title="Успех",
                        content=f"Поставщик '{name}' успешно создан",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
                    self.accept()
                else:
                    InfoBar.error(
                        title="Ошибка",
                        content=f"Не удалось создать поставщика '{name}'",
                        parent=self,
                        position=InfoBarPosition.TOP,
                        duration=3000
                    )
        except Exception as e:
            InfoBar.error(
                title="Ошибка",
                content=f"Не удалось сохранить поставщика: {str(e)}",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=5000
            )
        finally:
            db.close()


class MaterialRequestDialog(QDialog):
    def __init__(self, supplier_id, parent=None):
        super().__init__(parent)
        self.supplier_id = supplier_id
        
        # Set window title
        self.setWindowTitle("Формирование запроса материалов")
        
        # Initialize controllers
        self.provider_controller = ProviderController()
        self.material_controller = MaterialController()
        self.material_selections = {}
        
        # Setup UI
        self._setup_ui()
        
        # Load data
        self.load_data()
        
        # Set dialog size
        self.resize(700, 600)
        self.setStyleSheet("""
            QDialog {
                background-color: white;
            }
            QLineEdit {
                padding: 8px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: #f9f9f9;
            }
            QLineEdit[readOnly="true"] {
                background-color: #f0f0f0;
            }
            QLabel {
                font-size: 13px;
            }
            QCheckBox {
                padding: 5px;
                font-size: 13px;
            }
            QPushButton {
                padding: 8px 16px;
                font-weight: bold;
            }
        """)
        
    def _setup_ui(self):
        # Main layout
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(30, 30, 30, 30)
        
        # Title
        title_label = StrongBodyLabel("Формирование запроса материалов")
        title_label.setStyleSheet("font-size: 18px; margin-bottom: 10px;")
        main_layout.addWidget(title_label)
        
        # Supplier info section
        supplier_card = CardWidget()
        supplier_card.setStyleSheet("background-color: #f8f9fa; padding: 10px;")
        supplier_layout = QFormLayout(supplier_card)
        supplier_layout.setVerticalSpacing(15)
        supplier_layout.setLabelAlignment(Qt.AlignmentFlag.AlignRight)
        
        self.supplier_name_label = LineEdit()
        self.supplier_name_label.setReadOnly(True)
        supplier_layout.addRow(BodyLabel("Поставщик:"), self.supplier_name_label)
        
        self.supplier_contact_label = LineEdit()
        self.supplier_contact_label.setReadOnly(True)
        supplier_layout.addRow(BodyLabel("Контакт:"), self.supplier_contact_label)
        main_layout.addWidget(supplier_card)
        
        # Materials section
        materials_title = StrongBodyLabel("Выберите материалы для запроса:")
        materials_title.setStyleSheet("margin-top: 10px; margin-bottom: 5px;")
        main_layout.addWidget(materials_title)
        
        # Materials container
        materials_card = CardWidget()
        materials_card.setStyleSheet("background-color: #f8f9fa;")
        materials_inner_layout = QVBoxLayout(materials_card)
        
        # Materials scroll area
        materials_widget = QWidget()
        self.materials_layout = QVBoxLayout(materials_widget)
        self.materials_layout.setSpacing(15)
        
        materials_scroll = ScrollArea()
        materials_scroll.setWidget(materials_widget)
        materials_scroll.setWidgetResizable(True)
        materials_scroll.setMinimumHeight(250)
        materials_scroll.setStyleSheet("border: none; background-color: transparent;")
        materials_inner_layout.addWidget(materials_scroll)
        main_layout.addWidget(materials_card)
        
        # Comments
        comments_label = BodyLabel("Комментарий к запросу:")
        comments_label.setStyleSheet("margin-top: 10px;")
        main_layout.addWidget(comments_label)
        
        self.comments_edit = LineEdit()
        self.comments_edit.setPlaceholderText("Введите комментарий (необязательно)")
        self.comments_edit.setMinimumHeight(40)
        main_layout.addWidget(self.comments_edit)
        
        # Buttons
        buttons_layout = QHBoxLayout()
        buttons_layout.setContentsMargins(0, 20, 0, 0)
        buttons_layout.setSpacing(15)
        
        self.generate_button = PushButton("Сформировать запрос")
        self.generate_button.setIcon(FluentIcon.DOCUMENT)
        self.generate_button.setMinimumHeight(40)
        self.generate_button.setStyleSheet("background-color: #0078d4; color: white;")
        self.generate_button.clicked.connect(self.generate_document)
        buttons_layout.addWidget(self.generate_button)
        
        cancel_button = PushButton("Отмена")
        cancel_button.setMinimumHeight(40)
        cancel_button.clicked.connect(self.reject)
        buttons_layout.addWidget(cancel_button)
        main_layout.addLayout(buttons_layout)

    def load_data(self):
        """Load supplier and materials data"""
        try:
            db = SessionLocal()
            
            supplier = self.provider_controller.get_one(db, self.supplier_id)
            if not supplier:
                InfoBar.error(title="Ошибка", content="Поставщик не найден", parent=self)
                self.reject()
                return
                
            self.supplier_name_label.setText(supplier.name)
            
            contact_info = ""
            if supplier.phone: contact_info += f"Тел: {supplier.phone}"
            if supplier.mail:
                if contact_info: contact_info += ", "
                contact_info += f"Email: {supplier.mail}"
            self.supplier_contact_label.setText(contact_info if contact_info else "Нет контактной информации")
            
            materials = self.material_controller.get_all(db)
            # Clear previous items in materials_layout
            while self.materials_layout.count():
                child = self.materials_layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()
            
            if not materials:
                no_mat_label = BodyLabel("Нет доступных материалов", parent=self)
                no_mat_label.setStyleSheet("color: #666; padding: 20px; font-style: italic; text-align: center;")
                self.materials_layout.addWidget(no_mat_label)
                self.generate_button.setEnabled(False)
                return
                
            self.generate_button.setEnabled(True) # Enable if materials exist
            for material in materials:
                # Create a card for each material item
                material_widget = QWidget(self)
                material_widget.setStyleSheet("background-color: white; border-radius: 6px; padding: 5px;")
                
                mat_item_layout = QHBoxLayout(material_widget)
                mat_item_layout.setContentsMargins(10, 10, 10, 10)
                mat_item_layout.setSpacing(15)
                
                checkbox = QCheckBox(f"{material.type} (в наличии: {material.balance}, цена: {material.price} ₽)", parent=material_widget)
                checkbox.setChecked(False)
                checkbox.setStyleSheet("font-size: 14px;")
                mat_item_layout.addWidget(checkbox)
                
                quantity_spin = SpinBox(parent=material_widget)
                quantity_spin.setRange(1, material.balance)  # Set max to available balance
                quantity_spin.setValue(1)
                quantity_spin.setEnabled(False)
                quantity_spin.setFixedWidth(100)
                quantity_spin.setStyleSheet("QSpinBox { color: black; background-color: white; border: 1px solid #ccc; border-radius: 4px; padding: 4px; }")
                mat_item_layout.addWidget(quantity_spin)
                
                checkbox.stateChanged.connect(lambda state, spin=quantity_spin, mat_id=material.id: 
                    self.toggle_material_selection(state, spin, mat_id))
                    
                self.materials_layout.addWidget(material_widget)
        except Exception as e:
            InfoBar.error(title="Ошибка", content=f"Не удалось загрузить данные: {str(e)}", parent=self)
            self.reject()
        finally:
            db.close()
            
    def toggle_material_selection(self, state, spinbox, material_id):
        """Handle material selection"""
        spinbox.setEnabled(state == Qt.CheckState.Checked.value)
        
        if state == Qt.CheckState.Checked.value:
            self.material_selections[material_id] = spinbox.value()
        else:
            if material_id in self.material_selections:
                del self.material_selections[material_id]
                
        # Update spinbox value change handler
        if state == Qt.CheckState.Checked.value:
            spinbox.valueChanged.connect(
                lambda value, mat_id=material_id: self.update_material_quantity(mat_id, value)
            )
        
    def update_material_quantity(self, material_id, value):
        """Update quantity for selected material"""
        self.material_selections[material_id] = value
        
    def generate_document(self):
        """Generate Word document for material request"""
        if not self.material_selections:
            InfoBar.warning(
                title="Внимание",
                content="Выберите хотя бы один материал для запроса",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            return
            
        try:
            db = SessionLocal()
            
            # Get supplier
            supplier = self.provider_controller.get_one(db, self.supplier_id)
            if not supplier:
                InfoBar.error(
                    title="Ошибка",
                    content="Поставщик не найден",
                    parent=self,
                    position=InfoBarPosition.TOP,
                    duration=3000
                )
                return
                
            # Get selected materials
            materials_data = []
            for material_id, quantity in self.material_selections.items():
                material = self.material_controller.get_one(db, material_id)
                if material:
                    materials_data.append({
                        "material": material,
                        "quantity": quantity
                    })
            
            # Create Word document
            doc = docx.Document()
            
            # Document title
            title = doc.add_heading("ЗАПРОС МАТЕРИАЛОВ", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_paragraph = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Supplier info section
            doc.add_heading("Поставщик:", level=2)
            supplier_info = doc.add_paragraph()
            supplier_info.add_run(f"Название: {supplier.name}\n")
            supplier_info.add_run(f"ИНН: {supplier.inn}\n")
            
            if supplier.phone:
                supplier_info.add_run(f"Телефон: {supplier.phone}\n")
                
            if supplier.mail:
                supplier_info.add_run(f"Email: {supplier.mail}\n")
                
            if supplier.address:
                supplier_info.add_run(f"Адрес: {supplier.address}\n")
                
            doc.add_paragraph()
            
            # Requested materials section
            doc.add_heading("Запрашиваемые материалы:", level=2)
            
            # Create table for materials
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Set headers
            headers = table.rows[0].cells
            headers[0].text = "Материал"
            headers[1].text = "Количество"
            headers[2].text = "Примечание"
            
            # Add materials to table
            for item in materials_data:
                material = item["material"]
                quantity = item["quantity"]
                
                row = table.add_row().cells
                row[0].text = material.type
                row[1].text = str(quantity)
                row[2].text = ""
                
            doc.add_paragraph()
            
            # Comments
            comments = self.comments_edit.text().strip()
            if comments:
                doc.add_heading("Комментарии:", level=2)
                doc.add_paragraph(comments)
                doc.add_paragraph()
            
            # Signatures section
            doc.add_paragraph("Подпись заказчика: ____________________")
            doc.add_paragraph()
            doc.add_paragraph("Подпись поставщика: ____________________")
            
            # Save to temporary file
            fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            doc.save(temp_path)
            
            # Open the document with default application
            os.startfile(temp_path)
            
            InfoBar.success(
                title="Успех",
                content=f"Запрос материалов для поставщика '{supplier.name}' успешно сформирован",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=3000
            )
            
            self.accept()
            
        except Exception as e:
            InfoBar.error(
                title="Ошибка",
                content=f"Не удалось сформировать запрос: {str(e)}",
                parent=self,
                position=InfoBarPosition.TOP,
                duration=5000
            )
        finally:
            db.close() 