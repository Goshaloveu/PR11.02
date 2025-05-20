import os
import tempfile
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx
from openpyxl import Workbook


def get_documents_path():
    """Get path to save documents."""
    # Try to use user's Documents folder
    user_home = os.path.expanduser("~")
    user_docs = os.path.join(user_home, "Documents")
    terra_docs = os.path.join(user_docs, "Terra_Documents")
    
    # Create folder if it doesn't exist
    if not os.path.exists(terra_docs):
        try:
            os.makedirs(terra_docs)
        except Exception:
            # Fall back to temp directory
            terra_docs = os.path.join(tempfile.gettempdir(), "Terra_Documents")
            if not os.path.exists(terra_docs):
                os.makedirs(terra_docs)
                
    return terra_docs


def generate_order_statement(order_data):
    """Generate PDF statement for an order."""
    # Setup save path
    docs_path = get_documents_path()
    order_id = order_data.get('id_', 'unknown')
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"order_statement_{order_id}_{timestamp}.pdf"
    output_path = os.path.join(docs_path, filename)
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    elements = []
    
    # Set up styles
    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    subtitle_style = styles["Heading2"]
    normal_style = styles["Normal"]
    
    # Add title
    elements.append(Paragraph("Отчет по заказу", title_style))
    elements.append(Spacer(1, 12))
    
    # Order details
    elements.append(Paragraph(f"Заказ №: {order_data.get('id_', '')}", subtitle_style))
    elements.append(Spacer(1, 6))
    
    # Create a table for order info
    order_info = [
        ["Дата заказа:", str(order_data.get('date', ''))],
        ["Статус:", order_data.get('status', '')],
        ["Срок производства:", f"{order_data.get('prod_period', 'Не указан')} дней" if order_data.get('prod_period') else "Не указан"]
    ]
    
    # Add client info
    client = order_data.get('client', {})
    if client:
        client_name = f"{client.get('first', '')} {client.get('last', '')}"
        client_phone = client.get('phone', 'Не указан')
        client_email = client.get('mail', 'Не указан')
        
        order_info.extend([
            ["Клиент:", client_name],
            ["Телефон клиента:", client_phone],
            ["Email клиента:", client_email]
        ])
    
    # Add employee info
    worker = order_data.get('worker', {})
    if worker:
        worker_name = f"{worker.get('first', '')} {worker.get('last', '')}"
        worker_position = worker.get('position', '')
        
        order_info.extend([
            ["Сотрудник:", worker_name],
            ["Должность:", worker_position]
        ])
    else:
        order_info.append(["Сотрудник:", "Не назначен"])
    
    # Create table with info
    info_table = Table(order_info, colWidths=[150, 350])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
    ]))
    
    elements.append(info_table)
    elements.append(Spacer(1, 12))
    
    # Materials section
    elements.append(Paragraph("Материалы", subtitle_style))
    elements.append(Spacer(1, 6))
    
    # Create table headers for materials
    materials_data = [["Материал", "Количество", "Цена за ед.", "Стоимость"]]
    
    # Add materials data
    materials = order_data.get('materials_on_order', [])
    total_cost = 0
    
    for mat in materials:
        material = mat.get('material', {})
        material_type = material.get('type', '')
        amount = mat.get('amount', 0)
        price = material.get('price', 0)
        cost = amount * price
        total_cost += cost
        
        materials_data.append([
            material_type,
            str(amount),
            f"{price} ₽",
            f"{cost} ₽"
        ])
    
    # Add total row
    materials_data.append(["", "", "ИТОГО:", f"{total_cost} ₽"])
    
    # Create materials table
    materials_table = Table(materials_data, colWidths=[200, 100, 100, 100])
    materials_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 1), (3, -1), 'RIGHT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
    ]))
    
    elements.append(materials_table)
    elements.append(Spacer(1, 24))
    
    # Add signature section
    elements.append(Paragraph("Подписи сторон:", normal_style))
    elements.append(Spacer(1, 24))
    
    # Create signature table
    signature_data = [
        ["Клиент", "Сотрудник"],
        ["_______________________", "_______________________"],
        ["", ""],
        ["Дата: ________________", "Дата: ________________"]
    ]
    
    signature_table = Table(signature_data, colWidths=[250, 250])
    signature_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(signature_table)
    
    # Build PDF
    try:
        doc.build(elements)
        return output_path
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None


def generate_materials_report(materials_data):
    """Generate Excel report with materials."""
    # Setup save path
    docs_path = get_documents_path()
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"materials_report_{timestamp}.xlsx"
    output_path = os.path.join(docs_path, filename)
    
    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Материалы"
    
    # Add headers
    headers = ["ID", "Тип материала", "Остаток", "Цена"]
    for col, header in enumerate(headers, start=1):
        ws.cell(row=1, column=col, value=header)
    
    # Add materials data
    for row, material in enumerate(materials_data, start=2):
        ws.cell(row=row, column=1, value=material.get('id_', ''))
        ws.cell(row=row, column=2, value=material.get('type', ''))
        ws.cell(row=row, column=3, value=material.get('balance', 0))
        ws.cell(row=row, column=4, value=material.get('price', 0))
    
    # Style the worksheet
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[chr(64 + col)].width = 20
    
    # Add total row
    total_row = row + 1
    ws.cell(row=total_row, column=2, value="ИТОГО:")
    ws.cell(row=total_row, column=3, value=f"=SUM(C2:C{row})")
    
    # Save workbook
    try:
        wb.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error generating Excel report: {e}")
        return None


def generate_supplier_request(supplier_data, materials_data, comments=None):
    """Generate Word document for material request to supplier."""
    # Setup save path
    docs_path = get_documents_path()
    supplier_name = supplier_data.get('name', 'unknown').replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"supplier_request_{supplier_name}_{timestamp}.docx"
    output_path = os.path.join(docs_path, filename)
    
    # Create Word document
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Заявка на материалы', 0)
    
    # Add date
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    
    # Add supplier info
    doc.add_heading('Информация о поставщике', level=1)
    supplier_table = doc.add_table(rows=1, cols=2)
    supplier_table.style = 'Table Grid'
    
    # Add supplier headers
    header_cells = supplier_table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Значение'
    
    # Add supplier data
    supplier_info = [
        ['Наименование', supplier_data.get('name', '')],
        ['ИНН', supplier_data.get('inn', '')],
        ['Телефон', supplier_data.get('phone', '')],
        ['Email', supplier_data.get('mail', '')],
        ['Адрес', supplier_data.get('address', '')]
    ]
    
    for param, value in supplier_info:
        row_cells = supplier_table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value
    
    doc.add_paragraph()
    
    # Add materials section
    doc.add_heading('Запрашиваемые материалы', level=1)
    
    materials_table = doc.add_table(rows=1, cols=3)
    materials_table.style = 'Table Grid'
    
    # Add materials headers
    header_cells = materials_table.rows[0].cells
    header_cells[0].text = '№'
    header_cells[1].text = 'Материал'
    header_cells[2].text = 'Количество'
    
    # Add materials data
    for i, material in enumerate(materials_data, start=1):
        row_cells = materials_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = material.get('type', '')
        row_cells[2].text = str(material.get('amount', ''))
    
    doc.add_paragraph()
    
    # Add comments if provided
    if comments:
        doc.add_heading('Комментарии', level=1)
        doc.add_paragraph(comments)
        doc.add_paragraph()
    
    # Add signature section
    doc.add_heading('Подпись', level=1)
    doc.add_paragraph("__________________________ / ___________________________")
    doc.add_paragraph("       (подпись)                                 (расшифровка)")
    
    # Save document
    try:
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return None


def generate_order_receipt(order_data):
    """Generate DOCX receipt for client after order creation."""
    # Setup save path
    docs_path = get_documents_path()
    order_id = order_data.id if hasattr(order_data, 'id') else order_data.get('id', 'unknown')
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"order_receipt_{timestamp}.docx"
    output_path = os.path.join(docs_path, filename)
    
    # Create Word document
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Квитанция оплаты', 0)
    
    # Add date
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    
    # Add order info section
    doc.add_heading('Информация о заказе', level=1)
    order_table = doc.add_table(rows=1, cols=2)
    order_table.style = 'Table Grid'
    
    # Add order headers
    header_cells = order_table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Значение'
    
    # Add order data
    order_date = ""
    if hasattr(order_data, 'date'):
        order_date = order_data.date.strftime("%d.%m.%Y %H:%M") if order_data.date else ""
    else:
        order_date = order_data.get('date', '').strftime("%d.%m.%Y %H:%M") if order_data.get('date') else ""
    
    order_info = [
        ['Дата заказа:', order_date],
    ]
    
    # Add client info
    client = order_data.client if hasattr(order_data, 'client') else order_data.get('client', {})
    if client:
        client_name = f"{client.first if hasattr(client, 'first') else client.get('first', '')} {client.last if hasattr(client, 'last') else client.get('last', '')}"
        client_phone = client.phone if hasattr(client, 'phone') else client.get('phone', 'Не указан')
        
        order_info.extend([
            ['Клиент:', client_name],
            ['Телефон клиента:', client_phone],
        ])
    
    # Add table rows
    for param, value in order_info:
        row_cells = order_table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Materials section
    doc.add_heading('Материалы', level=1)
    
    materials_table = doc.add_table(rows=1, cols=4)
    materials_table.style = 'Table Grid'
    
    # Add materials headers
    header_cells = materials_table.rows[0].cells
    header_cells[0].text = 'Материал'
    header_cells[1].text = 'Количество'
    header_cells[2].text = 'Цена за ед.'
    header_cells[3].text = 'Стоимость'
    
    # Add materials data
    materials = []
    if hasattr(order_data, 'materials_on_order'):
        materials = order_data.materials_on_order
    elif hasattr(order_data, 'materials_link'):
        materials = order_data.materials_link
    elif isinstance(order_data, dict) and 'materials_on_order' in order_data:
        materials = order_data.get('materials_on_order', [])
    
    total_cost = 0
    
    for mat in materials:
        if hasattr(mat, 'material'):
            material = mat.material
            material_type = material.type if hasattr(material, 'type') else material.get('type', '')
            amount = mat.amount if hasattr(mat, 'amount') else mat.get('amount', 0)
            price = material.price if hasattr(material, 'price') else material.get('price', 0)
        else:
            material = mat.get('material', {})
            material_type = material.get('type', '')
            amount = mat.get('amount', 0)
            price = material.get('price', 0)
            
        cost = amount * price
        total_cost += cost
        
        row_cells = materials_table.add_row().cells
        row_cells[0].text = material_type  # Используем тип материала вместо ID
        row_cells[1].text = str(amount)
        row_cells[2].text = f"{price} ₽"
        row_cells[3].text = f"{cost} ₽"
    
    # Add total row
    row_cells = materials_table.add_row().cells
    row_cells[0].text = ""
    row_cells[1].text = ""
    row_cells[2].text = "ИТОГО:"
    row_cells[3].text = f"{total_cost} ₽"
    
    doc.add_paragraph()
    
    # Add thank you note
    doc.add_paragraph("Спасибо за выбор нашей компании!")
    doc.add_paragraph()
    
    # Add footer with current date
    doc.add_paragraph(f"Дата формирования квитанции: {datetime.now().strftime('%d.%m.%Y')}")
    
    # Save document
    try:
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error generating DOCX receipt: {e}")
        return None 